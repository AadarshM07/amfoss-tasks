import csv
from telegram import InlineKeyboardButton, InlineKeyboardMarkup,InputFile
from telegram.ext import Application, CommandHandler, CallbackQueryHandler,ContextTypes,MessageHandler,filters, CallbackContext
from telegram import Update, InlineQueryResultArticle, InputTextMessageContent
from telegram.ext import InlineQueryHandler
from uuid import uuid4
from telegram.helpers import escape_markdown
import os
from docx import Document
import logging
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)



user_state={}
L=""
async def start(update, context):
    await update.message.reply_text("Welcome to PagePal! I'm here to fetch book details for you. Use /help to see all commands and get started. Happy reading! 📚")

async def book(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    user_state[user_id] = 'book'  # Set the state for the user

    keyboard = [
        [InlineKeyboardButton("Drama", callback_data='1')],
        [InlineKeyboardButton("Business & Economics", callback_data='2')],
        [InlineKeyboardButton("Fiction", callback_data='3')],
        [InlineKeyboardButton("Literary Criticism", callback_data='4')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await update.message.reply_text("Please select a genre:", reply_markup=reply_markup)

async def button(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    user_id = query.from_user.id

    if user_id in user_state and user_state[user_id] == 'book':
        await query.answer()  # Acknowledge the callback query
        genre_map = {
            "1": "drama",
            "2": "business & economics",
            "3": "fiction",
            "4": "literary criticism"
        }
        genre = genre_map.get(query.data)

        if genre:
            found_books = []
            with open('book.csv', mode='r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                for k in reader:
                    if k[6].strip().lower() == genre:
                        found_books.append(f"Title: {k[0]}\nAuthor: {k[1]}\n")

            if found_books:
                await query.edit_message_text("\n\n".join(found_books))
            else:
                await query.edit_message_text("No books found for the selected genre.")
        else:
            await query.edit_message_text("Invalid choice, please select a valid number from 1 to 4.")

async def preview_command(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    user_state[user_id] = 'awaiting_book_name' 
    await update.message.reply_text("Please enter the name of the book you want to preview:")

async def handle_message(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    if user_id in user_state and user_state[user_id] == 'awaiting_book_name':
        book_name = update.message.text
        user_state[user_id] = 'completed'  

        # Check if the book is in the CSV file
        with open('book.csv', mode='r', newline='', encoding='utf-8') as file:
            reader = csv.reader(file)
            found = False
            for row in reader:
                if row[0].lower() == book_name.lower():  
                    preview_link = row[8] 
                    await update.message.reply_text(f"Title: {row[0]}\nAuthor: {row[1]}\nPreview: {preview_link}")
                    found = True
                    break
            if not found:
                await update.message.reply_text("No book found with that name.")

async def reading_list(update: Update, context: CallbackContext) -> None:
    user_id = update.message.from_user.id
    user_state[user_id] = 'read'
    key = [
        [
            InlineKeyboardButton("Add a book", callback_data='add_book'),
            InlineKeyboardButton("Delete a book", callback_data='delete_book'),
        ],
        [InlineKeyboardButton("View Reading List", callback_data='view_list')],
    ]
    reply_markup = InlineKeyboardMarkup(key)

    await update.message.reply_text("What would you like to do?", reply_markup=reply_markup)

async def handle_callback_query(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    user_id = query.from_user.id

    if user_id in user_state and user_state[user_id] == 'read':
        await query.answer()  # Acknowledge the callback query

        if query.data == 'add_book':
            context.user_data['action'] = 'A'
            user_state[user_id] = 'handle'
            await query.message.reply_text("Enter the book name you want to add:")

        elif query.data == 'delete_book':
            context.user_data['action'] = 'D'
            user_state[user_id] = 'handle'
            await query.message.reply_text("Enter the book name you want to delete:")

        elif query.data == 'view_list':
            if os.path.exists('readlist.docx'):
                with open('readlist.docx', 'rb') as f:
                    await context.bot.send_document(chat_id=query.message.chat_id, document=InputFile(f, 'readlist.docx'))
            else:
                await query.message.reply_text("No reading list found.")

async def handle_book_input(update: Update, context: CallbackContext) -> None:
    book_name = update.message.text
    action = context.user_data.get('action')
    user_id = update.message.from_user.id
    file_path = "readlist.docx"
    

    if user_id in user_state and user_state[user_id] == 'handle':
        if action == 'A':
            await update.message.reply_text(f"Book '{book_name}' added to the reading list.")

            if not os.path.exists(file_path):
                doc = Document()
            else:
                doc = Document(file_path)
            doc.add_paragraph(book_name)
            doc.save(file_path)

        elif action == 'D':
            await update.message.reply_text(f"Book '{book_name}' deleted from the reading list.")
            if os.path.exists(file_path):
                doc = Document(file_path)
                paragraphs = doc.paragraphs
                paragraphs_to_remove = [p for p in paragraphs if p.text == book_name]

                for paragraph in paragraphs_to_remove:
                    doc._element.body.remove(paragraph._element)

                doc.save(file_path)
            else:
                await update.message.reply_text("No reading list found.")

        # Reset state after processing
        user_state[user_id] = 'completed'
        context.user_data['action'] = None

async def help_command(update: Update, context: CallbackContext) -> None:
    await update.message.reply_text("""
        **The following commands are available:**

        /start: Get a welcome message.
        /book : Enter a genre, and I'll fetch you with book details.
        /preview : Get a preview link for the specified book.
        /reading_list: Displays three buttons:
               -Add a book
               -Delete a book
               -View Reading List
        /help: Show this list of commands."""
    )

# Token and application setup
token = "BOT_TOKEN"
application = Application.builder().token(token).build()

application.add_handler(CommandHandler('start', start))
application.add_handler(CommandHandler('help', help_command))
application.add_handler(CommandHandler('book', book))
application.add_handler(CommandHandler('preview', preview_command))
application.add_handler(CommandHandler('reading_list', reading_list))

# Callback query handlers
application.add_handler(CallbackQueryHandler(button, pattern='^[1-4]$'))
application.add_handler(CallbackQueryHandler(handle_callback_query, pattern='^(add_book|delete_book|view_list)$'))

# Message handlers
application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_input))

application.run_polling()
